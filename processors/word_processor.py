"""
Word处理器
处理Word文件，支持.docx格式（原生解析）和.doc格式（通过LibreOffice转换为.docx后解析）
"""

import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Dict, Any, List

from .base import BaseProcessor, ProcessingResult

# Word处理相关导入
try:
    import docx
    from docx.document import Document as DocxDocument
    from docx.oxml.ns import qn
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table, _Cell
    from docx.text.paragraph import Paragraph
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False

# .doc转换需要的LibreOffice可执行文件候选路径
SOFFICE_CANDIDATES = [
    "soffice",
    "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    "/usr/bin/soffice",
    "/opt/libreoffice/program/soffice",
]

# Word标题样式 -> Markdown标题级别
HEADING_STYLE_LEVELS = {f"Heading {i}": i for i in range(1, 10)}
HEADING_STYLE_LEVELS["Title"] = 1


class WordProcessor(BaseProcessor):
    """Word文档处理器"""

    def __init__(self, config: Dict[str, Any] = None):
        super().__init__(config)

        if not WORD_AVAILABLE:
            raise ImportError("Word处理相关库未安装(python-docx)，无法处理Word文件")

    def get_supported_extensions(self) -> list:
        """获取支持的文件扩展名"""
        return ['.doc', '.docx']

    def _find_soffice(self) -> str:
        """查找LibreOffice可执行文件路径"""
        for candidate in SOFFICE_CANDIDATES:
            resolved = shutil.which(candidate) if "/" not in candidate else candidate
            if resolved and Path(resolved).exists():
                return resolved
        return None

    def _convert_doc_to_docx(self, file_path: Path) -> Path:
        """
        使用LibreOffice将旧版.doc转换为.docx

        Args:
            file_path: .doc文件路径

        Returns:
            Path: 转换后的.docx临时文件路径

        Raises:
            RuntimeError: 未找到LibreOffice或转换失败
        """
        soffice = self._find_soffice()
        if not soffice:
            raise RuntimeError(
                "未找到LibreOffice(soffice)，无法转换.doc文件。"
                "请安装LibreOffice后重试。"
            )

        tmp_dir = tempfile.mkdtemp(prefix="doc2docx_")
        cmd = [
            soffice, "--headless", "--norestore",
            "--convert-to", "docx", "--outdir", tmp_dir, str(file_path)
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)

        converted_path = Path(tmp_dir) / (file_path.stem + ".docx")
        if result.returncode != 0 or not converted_path.exists():
            raise RuntimeError(
                f"LibreOffice转换.doc失败: {result.stderr.strip() or result.stdout.strip()}"
            )

        return converted_path

    def _iter_block_items(self, document: "DocxDocument"):
        """
        按原始顺序遍历文档正文中的段落和表格

        Args:
            document: python-docx Document对象

        Yields:
            Paragraph 或 Table 对象
        """
        body = document.element.body
        for child in body.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, document)
            elif isinstance(child, CT_Tbl):
                yield Table(child, document)

    def _paragraph_to_markdown(self, paragraph: "Paragraph") -> str:
        """将段落转换为Markdown文本"""
        text = paragraph.text.strip()
        if not text:
            return ""

        style_name = paragraph.style.name if paragraph.style else ""
        if style_name in HEADING_STYLE_LEVELS:
            level = min(HEADING_STYLE_LEVELS[style_name], 6)
            return f"{'#' * level} {text}"

        if style_name and ("List Bullet" in style_name or "List Paragraph" in style_name):
            return f"- {text}"
        if style_name and "List Number" in style_name:
            return f"1. {text}"

        return text

    def _table_to_markdown(self, table: "Table") -> str:
        """将表格转换为Markdown表格，处理水平合并单元格重复问题"""
        rows_data: List[List[str]] = []
        for row in table.rows:
            row_cells = []
            seen_tc_ids = set()
            for cell in row.cells:
                tc_id = id(cell._tc)
                if tc_id in seen_tc_ids:
                    continue
                seen_tc_ids.add(tc_id)
                cell_text = " ".join(p.text for p in cell.paragraphs).strip()
                row_cells.append(cell_text.replace("\n", " "))
            if row_cells:
                rows_data.append(row_cells)

        if not rows_data:
            return ""

        col_count = max(len(r) for r in rows_data)
        for r in rows_data:
            while len(r) < col_count:
                r.append("")

        header = "| " + " | ".join(rows_data[0]) + " |"
        separator = "| " + " | ".join(["---"] * col_count) + " |"
        body_lines = ["| " + " | ".join(r) + " |" for r in rows_data[1:]]

        return "\n".join([header, separator] + body_lines)

    def process(self, file_path: Path) -> ProcessingResult:
        """
        处理Word文件

        Args:
            file_path: Word文件路径

        Returns:
            ProcessingResult: 处理结果
        """
        converted_temp_path = None
        try:
            self.logger.info(f"开始处理Word文件: {file_path.name}")

            docx_path = file_path
            if file_path.suffix.lower() == '.doc':
                self.logger.info("检测到旧版.doc格式，使用LibreOffice转换为.docx")
                converted_temp_path = self._convert_doc_to_docx(file_path)
                docx_path = converted_temp_path

            document = docx.Document(str(docx_path))

            blocks = []
            table_count = 0
            for item in self._iter_block_items(document):
                if isinstance(item, Paragraph):
                    md = self._paragraph_to_markdown(item)
                    if md:
                        blocks.append(md)
                elif isinstance(item, Table):
                    md_table = self._table_to_markdown(item)
                    if md_table:
                        blocks.append(md_table)
                        table_count += 1

            markdown_content = f"# {file_path.stem}\n\n" + "\n\n".join(blocks)
            markdown_content = markdown_content.strip()

            if not markdown_content:
                return ProcessingResult(
                    success=False,
                    error="Word文件中没有提取到有效内容"
                )

            metadata = {
                'file_type': 'word',
                'file_extension': file_path.suffix.lower(),
                'paragraphs_count': len(document.paragraphs),
                'tables_count': table_count,
                'converted_from_doc': file_path.suffix.lower() == '.doc'
            }

            return ProcessingResult(
                success=True,
                content=markdown_content,
                metadata=metadata
            )

        except Exception as e:
            error_msg = f"处理Word文件失败: {str(e)}"
            self.logger.error(error_msg, exc_info=True)
            return ProcessingResult(
                success=False,
                error=error_msg
            )
        finally:
            if converted_temp_path and converted_temp_path.exists():
                self.cleanup_temp_files(converted_temp_path.parent)
